
from docx import Document
from docx.shared import Inches
import os

def build_report():
    doc = Document()
    doc.add_heading("A Novel Biomimetic Approach to Red Teaming A.I. Models", level=1)
    doc.add_paragraph("Author: Paul J. Lesevic")

    sections = [
        ("Executive Summary", "This report presents a biomimetic cognitive red teaming framework..."),
        ("1. Introduction", "The rise of large language models (LLMs)..."),
        ("2. Methodology", "We defined five cognitive profiles..."),
        ("2.1 Profile Simulation", "Profile details..."),
        ("2.2 Response Generation", "We used WebLLM..."),
        ("2.3 Annotation Process", "Each response was annotated..."),
        ("3. Results", "Empirical results and visualizations:"),
        ("3.1 Tag Distributions", "Auto-tagging undercounted hallucinations..."),
        ("3.2 Coherence Comparison", "Human evaluators consistently scored responses lower..."),
        ("3.3 Hallucination Heatmap", "Schizophrenia, Split-brain, and ADHD profiles..."),
        ("4. Interpretation", "Simulated cognitive distortions function as adversarial prompts..."),
        ("5. Limitations & Future Work", "Heuristics are blunt tools..."),
        ("6. Conclusion", "This work presents a red teaming blueprint grounded in human variability..."),
        ("Appendix A: Tagging Visuals", "See embedded figures above."),
        ("Appendix B: Full Entry-Level Tags & Justifications", "See supporting .docx with manual tags.")
    ]

    for title, content in sections:
        level = 2 if title[0].isdigit() or "Executive" in title or "Appendix" in title else 3
        doc.add_heading(title, level=level)
        doc.add_paragraph(content)

        # Placeholder for embedding figures
        img_map = {
            "3.1": "human_hallucination_heatmap.png",
            "3.2": "coherence_boxplot.png",
            "3.3": "human_contradiction_heatmap.png"
        }
        for key, img in img_map.items():
            if key in title and os.path.exists(f"data/{img}"):
                doc.add_picture(f"data/{img}", width=Inches(5.5))

    out_path = "data/LLM_Red_Team_Report.docx"
    doc.save(out_path)
    print("Report saved to:", out_path)

if __name__ == "__main__":
    build_report()
